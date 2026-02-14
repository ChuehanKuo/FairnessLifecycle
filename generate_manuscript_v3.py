"""Generate PLOS Digital Health manuscript v3 with lifecycle narrative."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    def h1(t):
        doc.add_heading(t, level=1)

    def h2(t):
        doc.add_heading(t, level=2)

    def p(t):
        doc.add_paragraph(t)

    def pc(t, sz=11, bold=False, italic=False):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        r.font.size = Pt(sz)
        r.bold = bold
        r.italic = italic

    # ================================================================
    # TITLE
    # ================================================================
    pc(
        "From benchmark to bedside: a fairness lifecycle evaluation of "
        "26 bias mitigation methods for health prediction in older Europeans",
        sz=15,
        bold=True,
    )
    doc.add_paragraph()
    pc("Short title: Fairness lifecycle in European health prediction AI", sz=10, italic=True)
    doc.add_paragraph()
    pc("[Author 1]\u00B9, [Author 2]\u00B2, [Author 3]\u00B9\u00B7\u00B3*")
    doc.add_paragraph()
    p("\u00B9 [Affiliation 1]")
    p("\u00B2 [Affiliation 2]")
    p("\u00B3 [Affiliation 3]")
    doc.add_paragraph()
    p("* Corresponding author: [email@institution.edu]")

    # ================================================================
    # ABSTRACT
    # ================================================================
    h1("Abstract")
    p(
        "Machine learning models are increasingly used to predict health outcomes in ageing "
        "populations, yet whether they perform equitably across income groups remains largely "
        "unexamined. Using data from 51,720 respondents across 27 European countries in the "
        "Survey of Health, Ageing and Retirement in Europe (SHARE) Wave 9, we quantify income-based "
        "disparities in prediction models for four health outcomes\u2014depression, life satisfaction, "
        "quality of life, and self-rated health\u2014and evaluate 26 bias mitigation methods spanning "
        "pre-processing, in-processing, post-processing, and data augmentation approaches. Baseline "
        "models exhibit true positive rate (TPR) gaps of 17\u201330 percentage points between the "
        "lowest- and highest-income quintiles, systematically under-detecting conditions in wealthier "
        "individuals. Rather than evaluating methods on benchmark metrics alone, we introduce a "
        "fairness lifecycle framework that traces each method through ten successive deployment "
        "stages: effectiveness, deployability under European data protection regulations, predictive "
        "accuracy preservation, cross-fold stability, computational feasibility, calibration "
        "integrity, threshold robustness, and regulatory compliance. This lifecycle evaluation "
        "reveals that 0 of 6 post-processing methods\u2014despite achieving the strongest benchmark "
        "performance\u2014survive deployment, because all require the protected attribute at inference "
        "time. Of 26 methods, only 3 survive all lifecycle stages: ExpGrad_TPR, ExpGrad_EO, and "
        "Reweighing. The first two achieve 86\u201394% reductions in income-based detection "
        "disparities at a cost of only 1\u20133 percentage points of AUROC, translating to "
        "approximately 1,723 additional correct identifications per screening wave in the "
        "highest-income quintile. Algorithmic fairness is not a property of a model; it is a "
        "property of a deployment lifecycle."
    )

    # ================================================================
    # AUTHOR SUMMARY
    # ================================================================
    h1("Author summary")
    p(
        "When AI models are used to screen older adults for depression, poor quality of life, or "
        "declining health, they can systematically miss cases in certain income groups. We analysed "
        "health data from over 51,000 Europeans aged 50 and older and tested 26 methods designed to "
        "make predictions fairer. Instead of simply comparing methods on benchmark numbers, we traced "
        "each one through the full journey from laboratory evaluation to real-world clinical "
        "deployment\u2014what we call the \u2018fairness lifecycle.\u2019 We discovered that the "
        "methods winning benchmarks (post-processing approaches) universally fail deployment because "
        "they require knowing each patient\u2019s income at every prediction, which European privacy "
        "law restricts. Of 26 methods, only 3 survived all deployment stages. The two strongest "
        "survivors reduce income-based detection gaps by 86\u201394% while barely affecting "
        "accuracy, potentially identifying approximately 1,723 additional cases per screening wave "
        "that would otherwise be missed. Our findings suggest that fairness research should evaluate "
        "methods not just on what they achieve in controlled settings, but on whether they can "
        "survive the full journey to the patient."
    )

    # ================================================================
    # INTRODUCTION
    # ================================================================
    h1("Introduction")
    p(
        "Europe faces a convergence of demographic ageing and persistent health inequality. More "
        "than 30% of Europeans will be over 65 by 2050, and the health burden of this transition "
        "falls disproportionately on the economically disadvantaged [2,3]. Among older Europeans, "
        "the lowest-income quintile experiences 1.7\u20132.6 times the prevalence of depression, "  # EUROD 1.65x, LS 2.15x, CASP 2.63x, SRH 1.90x
        "low life satisfaction, poor quality of life, and poor self-rated health compared with the "
        "highest-income quintile. These disparities are well documented but difficult to address at "
        "population scale using traditional screening methods."
    )
    p(
        "Machine learning (ML) models are increasingly deployed to identify individuals at risk of "
        "these conditions, enabling targeted interventions in population health screening programmes "
        "[5\u20137]. Yet a growing body of evidence demonstrates that predictive models can inherit "
        "and amplify the very disparities they are intended to help address. Obermeyer and colleagues "
        "showed that a widely used healthcare algorithm systematically underestimated the health "
        "needs of Black patients [1], catalysing urgent attention to algorithmic fairness in "
        "medicine. Growing evidence suggests that socioeconomic status (SES)-based disparities may "
        "be equally pernicious but remain critically understudied [10,11]."
    )
    p(
        "The algorithmic fairness literature has grown rapidly, but exhibits three blind spots that "
        "limit its clinical impact. First, existing work overwhelmingly focuses on racial and gender "
        "disparities; a scoping review of 91 clinical ML bias studies found that 64.8% evaluated "
        "race and 45.1% evaluated sex, while income and socioeconomic status were rarely examined "
        "[12,13]. Second, most fairness benchmarks evaluate methods on single outcomes or single "
        "datasets, making it difficult to assess whether findings generalise across the "
        "multidimensional nature of health [14,15]. Third, and most critically, evaluations "
        "overwhelmingly stop at the benchmark. Studies report fairness metric improvements in "
        "cross-validation and declare methods \u2018effective,\u2019 without examining whether "
        "those methods can survive the journey from benchmark to clinical deployment [16\u201319,63]."
    )
    p(
        "This third blind spot motivates our central conceptual contribution: the fairness "
        "lifecycle. We argue that algorithmic fairness is not a property of a model but a property "
        "of a deployment lifecycle. A method that achieves 90% reduction in true positive rate (TPR) "
        "gaps in cross-validation but requires the protected attribute at every prediction is not "
        "truly fair\u2014it is fair in the laboratory only. The journey from benchmark to bedside "
        "involves successive stages where fairness can be gained, degraded, or lost entirely: data "
        "collection, model training, intervention selection, accuracy\u2013fairness trade-off "
        "negotiation, stability testing across populations, computational feasibility assessment, "
        "calibration verification, threshold robustness analysis, regulatory review, and clinical "
        "deployment. We formalise this as a ten-stage lifecycle framework and apply it to the "
        "largest fairness benchmark in healthcare to date."
    )
    p(
        "The European regulatory landscape creates unique deployment constraints that make lifecycle "
        "evaluation particularly consequential. The General Data Protection Regulation (GDPR) "
        "restricts processing of sensitive personal data [25], while the EU AI Act creates narrow "
        "exceptions for bias detection in high-risk systems [24]. Whether and how fairness "
        "corrections can legally operate under these constraints depends critically on when the "
        "protected attribute is required\u2014a distinction invisible in benchmark-only evaluations "
        "but decisive for deployment [26\u201329]."
    )
    p(
        "This study addresses all three gaps through a unified investigation. We first quantify "
        "income-based disparities in health prediction models for older Europeans, establishing the "
        "scope and direction of the problem across four distinct health outcomes. We then evaluate "
        "26 bias mitigation methods\u2014the largest such benchmark in healthcare ML\u2014spanning "
        "pre-processing, in-processing, post-processing, and data augmentation categories. "
        "Crucially, we go beyond benchmark comparison by introducing and applying the fairness "
        "lifecycle framework, evaluating each method across ten deployment stages to identify which "
        "methods not only improve fairness metrics but can survive the full journey to clinical "
        "use. Finally, we translate surviving methods into clinical impact estimates, quantifying "
        "the number of additional patients who would be correctly identified in a screening "
        "programme."
    )

    # ================================================================
    # MATERIALS AND METHODS
    # ================================================================
    h1("Materials and methods")

    h2("Ethics statement")
    p(
        "This study used de-identified secondary data from SHARE Wave 9 (release 9.0.0, DOI: "
        "10.6103/SHARE.w9.900). SHARE data collection was approved by the ethics committee of "
        "the Max Planck Society and by ethics review boards in each participating country. All "
        "participants provided written informed consent. Our analysis of de-identified data was "
        "exempt from additional institutional review board approval."
    )

    h2("Study population and data source")
    p(
        "We used data from Wave 9 (2021\u20132022) of the Survey of Health, Ageing and Retirement "
        "in Europe (SHARE), a cross-national longitudinal study covering 27 European countries and "
        "Israel [4]. After excluding respondents with missing income data and applying quality "
        "filters, our sample comprised 51,720 respondents aged 50 and older. We used 175 predictor "
        "features spanning demographics, physical and mental health, cognitive function, social "
        "networks, housing conditions, and employment status."
    )

    h2("Health outcomes")
    p(
        "We predicted four binary outcomes representing distinct dimensions of wellbeing. EURO-D "
        "depression was defined as a score \u22654 on the 12-item scale [30] (prevalence: 26.3%). "
        "Life satisfaction was dichotomised at the median of a 0\u201310 scale (29.9%). Quality of "
        "life was measured using CASP-12 [31], dichotomised at the median (29.1%). Self-rated "
        "health was classified as poor/fair versus good or better [32] (36.3%). These four outcomes "
        "capture complementary dimensions of health and wellbeing and exhibit different "
        "socioeconomic gradients, enabling us to test whether fairness methods generalise across "
        "outcome types."
    )

    h2("Protected attribute")
    p(
        "Household income quintile (Q1 = lowest to Q5 = highest) served as the protected "
        "attribute, computed from total household income adjusted using the OECD-modified "
        "equivalence scale. Quintiles provided balanced group sizes (N = 10,123\u201310,500 per "
        "group) and captured the socioeconomic gradient in condition prevalence. Our primary "
        "fairness metric was the absolute difference in true positive rate (TPR) between Q1 and "
        "Q5. Secondary metrics included the equalised odds gap, demographic parity gap, and "
        "between-group AUROC gap."
    )

    h2("Baseline model and fairness evaluation")
    p(
        "All methods were compared against a LightGBM baseline [33], selected for computational "
        "efficiency and strong tabular data performance. Hyperparameters were tuned via Bayesian "
        "optimisation. The baseline included income quintile as an input feature but imposed no "
        "fairness constraints. Evaluation used five-fold stratified cross-validation preserving "
        "outcome and quintile proportions. Discrimination was assessed by AUROC, sensitivity, "
        "specificity, balanced accuracy, and Brier score. Directional consistency of fairness "
        "improvement across folds was assessed via sign tests."
    )

    h2("Bias mitigation methods")
    p(
        "We evaluated 26 methods in four categories using AI Fairness 360 [42], Fairlearn [43], "
        "and custom implementations. Pre-processing (n = 5): Reweighing [58], Disparate Impact "
        "Remover (DIR) [61], Learning Fair Representations (LFR) [60], Correlation Remover [43], "
        "and Sample Reweighting. In-processing (n = 13): Exponentiated Gradient with demographic "
        "parity, equalised odds, and TPR parity constraints (ExpGrad_DP, ExpGrad_EO, ExpGrad_TPR) "
        "[34], GridSearch with TPR and equalised odds constraints [34], Adversarial Debiasing [59], "
        "FairGBM [35], FairConstraints [36], Decoupled Classifier [39], GerryFair Classifier [40], "
        "Just Train Twice (JTT) [38], Group Distributionally Robust Optimisation (GroupDRO) [37], "
        "and Prejudice Remover [62]. Post-processing (n = 6): Threshold Optimisation for TPR "
        "parity, equalised odds, and demographic parity (ThreshOpt_TPR, ThreshOpt_EO, "
        "ThreshOpt_DP) [55], Equalised Odds Post-Processing (EqOdds_PP) [55], Reject Option "
        "Classification [41], and Group Calibration [46]. Data augmentation (n = 2): Conditional "
        "Tabular GAN (CTGAN) [45] and FairSMOTE [44]."
    )

    h2("Deployability assessment")
    p(
        "Given GDPR and EU AI Act constraints, we classified methods into three deployment tiers "
        "based on when the protected attribute is required. Tier 1 (14 methods): income needed "
        "only at training\u2014fully compatible with inference scenarios where income is "
        "unavailable. Tier 2 (6 methods): income needed at inference to set group-specific "
        "thresholds or apply post-hoc corrections. Tier 3 (6 methods): income used as a direct "
        "model input or required for feature transformation at inference. The EU AI Act\u2019s "
        "Article 10(5) creates a legal basis for processing sensitive data during training of "
        "high-risk AI systems [24], and health screening systems would likely qualify. Tier 1 "
        "methods are therefore fully GDPR/EU AI Act compatible, while Tier 2 and 3 methods face "
        "significant regulatory barriers."
    )

    h2("Fairness lifecycle evaluation framework")
    p(
        "We introduce a ten-stage lifecycle framework for evaluating the real-world viability of "
        "fairness methods beyond benchmark performance. The framework recognises that a method must "
        "survive multiple successive filters to deliver fairness impact in practice. The stages and "
        "their pass/fail criteria are as follows:"
    )
    p(
        "Stage 1 (Data characterisation): quantify baseline disparity across outcomes and income "
        "groups. Stage 2 (Baseline model): measure TPR gaps in the uncorrected model. Stage 3 "
        "(Effectiveness): the method must reduce disparity\u2014percentage reduction in TPR gap "
        "must be positive on at least three of four outcomes. Stage 4 (Accuracy preservation): "
        "AUROC must not drop more than 3 percentage points from baseline on any outcome, "
        "reflecting the threshold beyond which clinical utility is compromised. Stage 5 "
        "(Stability): the method must improve fairness in at least four of five cross-validation "
        "folds for all four outcomes (sign test \u2265 4/5). Stage 6 (Computational feasibility): "
        "training time must be practical for periodic retraining (<30 minutes per outcome). "
        "Stage 7 (Calibration): Brier score increase must remain below 0.03 from baseline. "
        "Stage 8 (Threshold robustness): fairness improvement must not be an artifact of a "
        "specific classification threshold. Stage 9 (Regulatory compliance): the method must not "
        "require the protected attribute at inference time (Tier 1 deployability). Stage 10 "
        "(Clinical deployment): surviving methods are translated into patient-level impact "
        "estimates."
    )
    p(
        "Methods are progressively filtered through stages 3\u20139. Only methods passing all "
        "stages are recommended for deployment and proceed to clinical impact estimation at "
        "stage 10."
    )

    h2("Clinical impact estimation")
    p(
        "For surviving methods, we translated TPR gap reductions into estimated additional correct "
        "identifications per screening wave using the formula: additional identifications = "
        "(baseline TPR gap \u2212 method TPR gap) \u00d7 N\u2085(positive), where N\u2085(positive) "
        "is the number of Q5 individuals with the condition, calculated from Q5 sample size "
        "(N = 10,123) and Q5-specific outcome prevalence rates."
    )

    h2("Robustness analyses")
    p(
        "Multi-threshold robustness analyses evaluated the baseline model under four threshold "
        "strategies: balanced accuracy optimisation (default), sensitivity \u226570%, "
        "prevalence-matched, and fixed 0.5. This analysis confirmed whether observed disparities "
        "were an artifact of threshold choice or reflected deeper model-level bias."
    )

    # ================================================================
    # RESULTS
    # ================================================================
    h1("Results")

    h2("Standard models exhibit substantial income-based disparities")
    p(
        "Before evaluating corrections, we characterised the scope of the fairness problem. A "
        "steep socioeconomic gradient was evident across all four health outcomes (Table 1): the "
        "lowest-income quintile (Q1) had 1.7\u20132.6 times the condition prevalence of the "
        "highest-income quintile (Q5). The baseline LightGBM model achieved strong overall "
        "discrimination (AUROC: EUROD 0.830, LS 0.805, CASP 0.883, SRH 0.893) but exhibited "
        "substantial TPR gaps between income groups: 0.167 (16.7 percentage points) for "
        "depression, 0.301 (30.1 pp) for life satisfaction, 0.280 (28.0 pp) for quality of "
        "life, and 0.181 (18.1 pp) for self-rated health."
    )
    p(
        "This pattern\u2014systematic under-detection in the highest-income group\u2014was "
        "consistent across all outcomes, all cross-validation folds, and all threshold strategies "
        "(Table S2). Multi-threshold analysis confirmed that baseline TPR gaps persisted regardless "
        "of operating point: gaps ranged from 0.193 to 0.217 for depression, 0.280 to 0.323 for "
        "life satisfaction, 0.291 to 0.307 for quality of life, and 0.205 to 0.231 for self-rated "
        "health across four threshold strategies. The disparity is not a threshold artifact; it is "
        "encoded in the model\u2019s learned representations."
    )
    p(
        "The likely mechanism is that wealthier individuals who do experience depression or poor "
        "quality of life present with feature profiles that differ from the typical affected "
        "population, which is predominantly low-income. The model learns \u2018low income predicts "
        "poor health\u2019 as a strong signal but struggles with the minority pattern of high "
        "income combined with poor health outcomes."
    )

    h2("Most bias mitigation methods reduce disparities, but effectiveness varies widely")
    p(
        "We evaluated 26 methods spanning four categories (Table 2). Effectiveness ranged from "
        "counterproductive (GerryFairClassifier increased the TPR gap by 24.6% on depression) to "
        "near-complete correction (ExpGrad_EO achieved 94.4% TPR gap reduction on life "
        "satisfaction). Five methods appeared consistently among the top performers across all "
        "four health outcomes: ExpGrad_TPR, ExpGrad_EO, FairConstraints, and two threshold "
        "optimisation variants (ThreshOpt_TPR, ThreshOpt_EO). However, several methods were "
        "counterproductive on at least one outcome: GerryFairClassifier, GroupCalibration, "
        "SampleReweighting, CTGAN, and DecoupledClassifier all increased disparities in at least "
        "one outcome, with GerryFairClassifier increasing the depression TPR gap by 24.6% and "
        "GroupCalibration increasing the life satisfaction gap by 11.1%."
    )

    h2("In-processing constraint methods achieve the best fairness\u2013accuracy balance")
    p(
        "The top three methods\u2014ExpGrad_TPR, ExpGrad_EO, and FairConstraints\u2014are all "
        "in-processing approaches that embed fairness constraints directly into the learning "
        "objective [34,36]. ExpGrad_TPR achieved 87\u201394% TPR gap reductions across outcomes "
        "with AUROC losses of only 0.8\u20132.8 percentage points. These methods work by "
        "iteratively re-weighting training examples to equalise group-level error rates, producing "
        "models that are intrinsically fairer without requiring post-hoc adjustments."
    )
    p(
        "Post-processing methods (ThreshOpt_TPR, ThreshOpt_EO, RejectOption) offered a distinct "
        "trade-off profile: 77\u201391% gap reductions at zero AUROC cost, because they adjust "
        "decision thresholds rather than the model itself. RejectOption was the single best method "
        "for self-rated health (90.5% reduction) and quality of life (90.0%). However, as we show "
        "below, this benchmark advantage is illusory from a deployment perspective."
    )
    p(
        "Pre-processing methods produced only modest improvements (29\u201343% for Reweighing, "
        "the best in this category), while data augmentation via CTGAN and FairSMOTE was largely "
        "ineffective (\u22123% to 16% reductions), consistent with evidence that naive synthetic "
        "data can amplify rather than correct bias [47]."
    )

    h2("The deployability cliff: post-processing methods fail deployment")
    p(
        "Our deployability assessment revealed a striking pattern that we term the "
        "\u2018deployability cliff.\u2019 All 6 post-processing methods require the protected "
        "attribute (income quintile) at inference time\u2014either to set group-specific decision "
        "thresholds or to apply group-conditional corrections. In a real screening programme, this "
        "means every patient would need to disclose income data at the point of prediction, the "
        "system would need to maintain a live income classification pipeline, and GDPR Article 9 "
        "restrictions on sensitive data processing would apply to every inference call."
    )
    p(
        "The deployability breakdown by category is as follows. Pre-processing: 2 of 5 methods "
        "(40%) are deployable. In-processing: 9 of 13 methods (69%) are deployable. "
        "Post-processing: 0 of 6 methods (0%) are deployable. Augmentation: 2 of 2 methods "
        "(100%) are deployable. The methods that achieve the strongest fairness improvements in "
        "controlled benchmarks are precisely the methods that cannot survive real-world deployment "
        "constraints. A benchmark-only evaluation would recommend post-processing methods as top "
        "performers; a lifecycle evaluation correctly identifies them as non-viable."
    )

    h2("The lifecycle filter: from 26 methods to 3 survivors")
    p(
        "Applying the fairness lifecycle framework progressively eliminates methods that fail "
        "real-world deployment requirements (Fig 4). Starting from 26 methods:"
    )
    p(
        "Stage 3 (Effectiveness): 4 methods are eliminated for increasing disparity on at least "
        "two of four outcomes: GerryFairClassifier (increased disparity on 3 outcomes), "
        "GroupCalibration (3 outcomes), CTGAN (3 outcomes), and GroupDRO (2 outcomes). "
        "Remaining: 22 methods."
    )
    p(
        "Stage 9 (Deployability): 11 methods are eliminated for requiring the protected attribute "
        "at inference, including all 6 post-processing methods plus AdvDebiasing, "
        "PrejudiceRemover, DecoupledClassifier, DIR, CorrelationRemover, and LFR. "
        "Remaining: 11 methods."
    )
    p(
        "Stage 4 (Accuracy): 5 methods are eliminated for exceeding the 3 percentage point AUROC "
        "threshold on at least one outcome: GridSearch_TPR (7.1 pp drop on LS), GridSearch_EO "
        "(8.9 pp), ExpGrad_DP (4.9 pp on LS), FairConstraints (5.0 pp on LS), and JTT (16.8 pp "
        "on EUROD). Remaining: 6 methods (ExpGrad_TPR, ExpGrad_EO, Reweighing, FairGBM, "
        "FairSMOTE, SampleReweighting)."
    )
    p(
        "Stages 5\u20137 (Stability, computational feasibility, calibration): all 6 remaining "
        "methods pass these criteria. All achieve sign test \u22654/5 across all outcomes, train "
        "in under 3 minutes per outcome, and maintain Brier score increases below 0.01."
    )
    p(
        "Stage 8 (Fairness threshold): 3 methods are eliminated for insufficient fairness "
        "improvement across all outcomes. FairGBM achieves only 8.8\u201317.3% TPR gap reduction, "
        "FairSMOTE achieves 11.0\u201316.1%, and SampleReweighting shows negative reductions on "
        "EUROD. All fall below our 25% minimum threshold for "
        "meaningful clinical impact on at least one outcome. Remaining: 3 methods."
    )
    p(
        "The three lifecycle survivors are ExpGrad_TPR (86.6\u201393.8% TPR gap reduction, "
        "0.8\u20132.8 pp AUROC cost), ExpGrad_EO (86.0\u201394.4% reduction, 1.1\u20133.0 pp "
        "cost), and Reweighing (28.7\u201342.9% reduction, <0.2 pp cost). ExpGrad_TPR and "
        "ExpGrad_EO are the clear leaders in fairness improvement, while Reweighing offers a "
        "minimal-intervention option that preserves full predictive performance at the cost of "
        "more modest fairness gains."
    )

    h2("Clinical impact of surviving methods")
    p(
        "For the top two survivors (ExpGrad_TPR and ExpGrad_EO), we translated TPR gap reductions "
        "into estimated additional correct identifications in the highest-income quintile "
        "(Q5, N = 10,123) per screening wave (Fig 5). For depression: baseline TPR gap of 0.167 "
        "reduced to 0.019 by ExpGrad_TPR, yielding approximately 315 additional correct "
        "identifications among Q5 individuals with depression. For life satisfaction: gap reduced "
        "from 0.301 to 0.019, yielding approximately 565 additional identifications. For quality "
        "of life: gap reduced from 0.280 to 0.021, yielding approximately 441 additional "
        "identifications. For self-rated health: gap reduced from 0.181 to 0.024, yielding "
        "approximately 402 additional identifications."
    )
    p(
        "Combined across all four outcomes, ExpGrad_TPR yields approximately 1,723 additional "
        "correct identifications per screening wave in Q5 alone, with proportional improvements "
        "expected across Q2\u2013Q4. In a pan-European screening programme serving the full SHARE "
        "population, this represents a substantial reduction in income-based health detection "
        "inequity."
    )

    h2("Results are robust across folds and threshold choices")
    p(
        "The top methods showed strong stability across cross-validation folds (Fig 1). ExpGrad_TPR "
        "and ExpGrad_EO achieved significant improvement in all five folds (5/5 sign test) for all "
        "four outcomes, with fold-level TPR gap standard deviations of 0.010\u20130.029. "
        "Multi-threshold analysis confirmed that the baseline\u2019s unfairness persists across all "
        "four threshold strategies, validating that in-processing corrections address root-cause "
        "model bias rather than threshold artifacts. Calibration degradation was minimal: Brier "
        "score increases of 0.003\u20130.010 for ExpGrad_TPR, well within clinically negligible "
        "range."
    )

    # ================================================================
    # DISCUSSION
    # ================================================================
    h1("Discussion")

    h2("Income bias in health AI: scope and mechanism")
    p(
        "Our first contribution is documenting the scope of income-based disparities in health "
        "prediction for older Europeans. The finding that standard models systematically "
        "under-detect conditions in the highest-income quintile may seem counterintuitive\u2014"
        "clinical attention typically focuses on disadvantaged populations. However, this pattern "
        "reflects an algorithmic mechanism, not a clinical priority: models trained predominantly "
        "on low-income positive cases learn feature patterns characteristic of that group and "
        "struggle to recognise the same conditions when they present with different feature "
        "profiles in wealthier individuals. The result is inequitable detection, regardless of "
        "direction."
    )
    p(
        "Notably, the magnitude of disparity varies substantially across outcomes. Life "
        "satisfaction shows the largest TPR gap (30.1 pp), likely because satisfaction is most "
        "strongly confounded by income\u2014the features that predict dissatisfaction in low-income "
        "individuals differ most from those in high-income individuals. Self-rated health shows "
        "the smallest gap (18.1 pp), perhaps because physical health indicators are less "
        "income-dependent. This outcome dependency underscores the importance of multi-outcome "
        "evaluation; a single-outcome study would capture only one dimension of the fairness "
        "problem."
    )

    h2("The fairness\u2013performance trade-off is navigable")
    p(
        "Our benchmark demonstrates that the fairness\u2013performance trade-off, while real, is "
        "far more manageable than previous healthcare fairness studies suggested. Pfohl and "
        "colleagues found near-universal performance degradation with fairness penalties on "
        "clinical risk prediction [15], and MEDFAIR concluded that no method consistently "
        "outperformed others across medical imaging tasks [14]. In contrast, we find that "
        "ExpGrad_TPR and ExpGrad_EO achieve 86\u201394% fairness improvement at a cost of only "
        "1\u20133 percentage points of AUROC. This is consistent with Rodolfa and colleagues\u2019 "
        "observation of \u2018negligible\u2019 fairness\u2013accuracy trade-offs in public policy "
        "applications [20]."
    )
    p(
        "The favourable trade-off likely reflects our specific methodological choices: gradient "
        "boosting on tabular survey data with a well-defined protected attribute. The causal "
        "fairness literature notes that from a causal perspective, some tension between fairness "
        "and utility always exists [21]. Our finding of small trade-offs reflects specific metric "
        "and data choices rather than a universal resolution. Whether a 1\u20133 percentage point "
        "AUROC reduction is acceptable depends on clinical context, and we leave this determination "
        "to the health systems that would deploy such models."
    )

    h2("The fairness lifecycle: a new evaluation paradigm")
    p(
        "This study\u2019s central conceptual contribution is the fairness lifecycle framework. "
        "Most fairness studies end at the benchmark: they report TPR gap reductions, equalised "
        "odds improvements, or demographic parity scores in cross-validation and declare methods "
        "\u2018effective.\u2019 But effectiveness in a benchmark is necessary, not sufficient, for "
        "real-world impact. The journey from benchmark to bedside involves multiple stages where "
        "fairness can degrade or fail entirely, and our lifecycle framework makes these failure "
        "modes explicit."
    )
    p(
        "Our lifecycle evaluation reveals three structural findings that benchmark-only studies "
        "cannot detect. First, the \u2018deployability cliff\u2019: all 6 post-processing methods "
        "require the protected attribute at inference, making them incompatible with European data "
        "protection requirements in most screening scenarios. This is particularly striking because "
        "post-processing methods often achieve the strongest benchmark performance (up to 91% TPR "
        "gap reduction for RejectOption on self-rated health). A benchmark study would recommend "
        "these methods; a lifecycle study correctly identifies them as non-viable."
    )
    p(
        "Second, the \u2018stability filter\u2019: several methods that show strong mean "
        "improvements fail to consistently improve fairness across population subsets. GroupDRO, "
        "for instance, preserves the highest AUROC (0.831 for depression) but improves fairness in "
        "only 1 of 5 cross-validation folds\u2014a method that works in 20% of populations and "
        "does nothing (or harm) in the other 80% would be rejected by any regulator. Third, the "
        "\u2018accuracy gate\u2019: aggressive constraint methods (GridSearch, ExpGrad_DP) achieve "
        "near-complete fairness correction but at AUROC costs of 5\u20139 percentage points that "
        "would compromise clinical utility."
    )
    p(
        "The progressive filtering from 26 methods to 3 survivors (with 2 achieving >85% "
        "reduction) illustrates why lifecycle evaluation matters. We propose that future fairness "
        "evaluations adopt lifecycle-oriented assessment as standard practice, reporting not just "
        "metric improvements but deployability constraints, stability across populations, "
        "calibration impact, and computational feasibility. The lifecycle framework is applicable "
        "beyond our specific context to any domain where fairness methods must survive real-world "
        "deployment [63,65]."
    )

    h2("From benchmarks to deployment: regulatory alignment")
    p(
        "Perhaps our most practically significant finding is that the best-performing deployable "
        "methods also face the fewest deployment barriers under European regulations. This "
        "alignment is not coincidental: in-processing constraint methods learn intrinsically "
        "fairer models by embedding fairness into model parameters during training, requiring "
        "the protected attribute only at training time [34]. The EU AI Act\u2019s Article 10(5) "
        "creates a legal basis for processing sensitive data during training of high-risk AI "
        "systems [24], and health screening systems would likely qualify."
    )
    p(
        "The practical advantage of income over race or health status as a protected attribute "
        "is significant: fewer regulatory barriers exist for income data, which is routinely "
        "collected in social surveys and administrative records. A health system could train a "
        "fairness-corrected model using linked income data, then deploy the trained model without "
        "requiring income information at the point of care. This separation of training-time and "
        "inference-time data requirements is precisely what makes Tier 1 methods deployable [27]."
    )

    h2("Clinical impact and public health implications")
    p(
        "Translating benchmark metrics into patient-level impact is essential for communicating "
        "the value of fairness corrections to policymakers and clinicians who do not engage with "
        "TPR gaps or AUROC scores. Our estimate of approximately 1,723 additional correct "
        "identifications per screening wave in Q5 alone makes the abstract concrete: these are "
        "individuals with depression, low quality of life, or poor health who would be missed by "
        "an unfair model but correctly identified by a fair one. Each represents a potential "
        "referral to appropriate care."
    )
    p(
        "The surviving methods (ExpGrad_TPR, ExpGrad_EO) are computationally tractable, training "
        "in approximately 130\u2013165 seconds per outcome. This enables quarterly retraining as "
        "new SHARE waves arrive, ensuring that fairness corrections remain calibrated to evolving "
        "population demographics. The third survivor, Reweighing, offers an even lighter-weight "
        "option (approximately 2 seconds per outcome) for systems that prioritise minimal "
        "computational overhead, albeit with more modest fairness gains (29\u201343% reduction)."
    )

    h2("Comparison with related work")
    p(
        "The most directly relevant prior work is Shen and colleagues\u2019 audit of SHARE-based "
        "models across 16 European countries [8], which found demographic parity ranging from 0.14 "
        "(Greece) to 0.96 (Switzerland) but applied no corrective methods. Our study complements "
        "their diagnostic work with systematic correction and lifecycle evaluation. Pfohl and "
        "colleagues [15] evaluated fairness methods for clinical risk prediction but focused on "
        "race with fewer methods and no deployment analysis. MEDFAIR [14] benchmarked fairness in "
        "medical imaging\u2014a different modality with different fairness dynamics. Chen and "
        "colleagues [16] conducted a comprehensive bias mitigation benchmark but without healthcare "
        "application or deployment framework. The FFB and ABCFair benchmarks [18,19] advanced "
        "methodological comparison but did not consider deployment constraints. Our lifecycle "
        "framework addresses the gap that all these prior studies share: the transition from "
        "benchmark evaluation to real-world deployment [64]."
    )

    h2("Limitations and future directions")
    p(
        "Several limitations should be acknowledged. First, we evaluate methods on a single "
        "dataset (SHARE Wave 9); cross-dataset and cross-country validation would strengthen "
        "generalisability, particularly given evidence that fairness corrections may not transfer "
        "across populations [54]. Second, income quintiles are a simplified operationalisation of "
        "socioeconomic status; wealth, education, and occupation may contribute independently to "
        "prediction disparities. Third, our analysis focused on binary group fairness (Q1 vs Q5); "
        "intersectional approaches that consider combinations of protected attributes may reveal "
        "additional disparities. Fourth, the lifecycle framework\u2019s stage-specific criteria "
        "(e.g., 3 pp AUROC threshold, 25% fairness minimum) involve judgment calls that other "
        "researchers may calibrate differently. Fifth, we did not evaluate causal fairness "
        "approaches [21], which may offer complementary insights. Future work should validate our "
        "lifecycle framework longitudinally across SHARE waves, extend to intersectional analysis, "
        "and integrate with causal fairness frameworks to distinguish between \u2018appropriate\u2019 "
        "and \u2018inappropriate\u2019 uses of income in health prediction."
    )

    # ================================================================
    # CONCLUSIONS
    # ================================================================
    h1("Conclusions")
    p(
        "Standard ML models for health prediction in older Europeans harbour substantial "
        "income-based disparities, systematically under-detecting conditions in the highest-income "
        "group by 17\u201330 percentage points in TPR. Through a fairness lifecycle evaluation of "
        "26 bias mitigation methods across four health outcomes, we demonstrate that the majority "
        "of methods achieving strong benchmark performance fail real-world deployment requirements. "
        "All post-processing methods require the protected attribute at inference, creating a "
        "fundamental incompatibility with European data protection regulations\u2014a finding we "
        "term the \u2018deployability cliff.\u2019 Of 26 methods, 3 survive the full lifecycle: "
        "ExpGrad_TPR, ExpGrad_EO, and Reweighing. The first two achieve 86\u201394% reductions "
        "in income-based detection disparities at 1\u20133 percentage point AUROC cost, "
        "translating to approximately 1,723 additional correct identifications per screening wave "
        "in the highest-income quintile. Algorithmic fairness is not a property of a model "
        "measured at the benchmark; it is a property of a deployment lifecycle that must be "
        "evaluated end-to-end."
    )

    # ================================================================
    # ACKNOWLEDGMENTS
    # ================================================================
    h1("Acknowledgments")
    p(
        "This paper uses data from SHARE Wave 9 (DOI: 10.6103/SHARE.w9.900). The SHARE data "
        "collection has been funded by the European Commission through FP5 (QLK6-CT-2001-00360), "
        "FP6 (SHARE-I3: RII-CT-2006-062193, COMPARE: CIT5-CT-2005-028857, SHARELIFE: "
        "CIT4-CT-2006-028812), FP7 (SHARE-PREP: GA N\u00b0211909, SHARE-LEAP: GA N\u00b0227822, "
        "SHARE M4: GA N\u00b0261982, DASISH: GA N\u00b0283646), Horizon 2020 (SHARE-DEV3: GA "
        "N\u00b0676536, SHARE-COHESION: GA N\u00b0870628, SERISS: GA N\u00b0654221, SSHOC: GA "
        "N\u00b0823782, SHARE-COVID19: GA N\u00b0101015924) and by DG Employment, Social Affairs "
        "& Inclusion, the US National Institute on Aging (U01_AG09740-13S2, P01_AG005842, "
        "P01_AG08291, P30_AG12815, R21_AG025169, Y1-AG-4553-01, IAG_BSR06-11, "
        "OGHA_04-064, HHSN271201300071C, RAG052527A), and various national funding sources."
    )

    # ================================================================
    # REFERENCES
    # ================================================================
    h1("References")
    refs = [
        "1. Obermeyer Z, Powers B, Vogeli C, Mullainathan S. Dissecting racial bias in an algorithm used to manage the health of populations. Science. 2019;366(6464):447-453.",
        "2. OECD. Health at a Glance: Europe 2024. Paris: OECD Publishing; 2024.",
        "3. World Health Organization. World report on ageing and health. Geneva: WHO; 2015.",
        "4. B\u00f6rsch-Supan A, Brandt M, Hunkler C, Kneip T, Korbmacher J, Malter F, et al. Data resource profile: The Survey of Health, Ageing and Retirement in Europe (SHARE). Int J Epidemiol. 2013;42(4):992-1001.",
        "5. Handing EP, Strobl C, Gerstorf D. Predictors of depression among middle-aged and older men and women in Europe: A machine learning approach. Lancet Reg Health Eur. 2022;18:100391.",
        "6. Prati G. Health status prediction for the elderly based on machine learning. Arch Gerontol Geriatr. 2022;101:104690.",
        "7. Lu Y, Yin J, Chen Z, Gong Y, Liu Z, Zhang L, et al. Determinants of depressive symptoms in multinational middle-aged and older adults. npj Digit Med. 2025;8:186.",
        "8. Shen X, Peng X, Zhang H, Wang J, Li K. Algorithm fairness in predicting unmet preventive care: Evidence from 16 European countries using SHARE. medRxiv [Preprint]. 2025.",
        "9. Rajkomar A, Hardt M, Howell MD, Corrado G, Chin MH. Ensuring fairness in machine learning to advance health equity. Ann Intern Med. 2018;169(12):866-872.",
        "10. Juhn YJ, Ryu E, Wi CI, King KS, Malik M, Juhn HJ, et al. Assessing socioeconomic bias in machine learning algorithms in health care: A case study of the HOUSES index. J Am Med Inform Assoc. 2022;29(7):1142-1151.",
        "11. Ciernikova S, Pospisilova E, Pec M, Bertinyova S, Nagyova I. Fairness and bias correction in machine learning for depression prediction across four study populations. Sci Rep. 2024;14:8567.",
        "12. Stahl D, Karcher N, Gkika S. Strategies to mitigate age-related bias in machine learning: Scoping review. JMIR Aging. 2024;7:e53564.",
        "13. Mehrabi N, Morstatter F, Saxena N, Leskovec J, Galstyan A. A survey on bias and fairness in machine learning. ACM Comput Surv. 2021;54(6):1-35.",
        "14. Zong Y, Yang Y, Hospedales T. MEDFAIR: Benchmarking fairness for medical imaging. In: ICLR; 2023.",
        "15. Pfohl SR, Foryciarz A, Shah NH. An empirical characterization of fair machine learning for clinical risk prediction. J Biomed Inform. 2021;113:103621.",
        "16. Chen Z, Zhang JM, Sarro F, Harman M. A comprehensive empirical study of bias mitigation methods for machine learning classifiers. ACM Trans Softw Eng Methodol. 2023;32(4):1-30.",
        "17. Friedler SA, Scheidegger C, Venkatasubramanian S, Choudhary S, Hamilton EP, Roth D. A comparative study of fairness-enhancing interventions in machine learning. In: FAccT; 2019. pp. 329-338.",
        "18. Han X, Tsang I, Zhu J, Liu T, Gong M, Liu F, et al. FFB: A fair fairness benchmark for in-processing group fairness methods. In: ICLR; 2024.",
        "19. Defrance M, Vanderdonckt J, Verleysen M. ABCFair: An adaptable benchmark approach for comparing fairness methods. In: NeurIPS Datasets and Benchmarks Track; 2024.",
        "20. Rodolfa KT, Lamba H, Ghani R. Empirical observation of negligible fairness-accuracy trade-offs in machine learning for public policy. Nat Mach Intell. 2021;3(10):896-904.",
        "21. Corbett-Davies S, Goel S. The measure and mismeasure of fairness: A critical review of fair machine learning. J Mach Learn Res. 2023;24(312):1-117.",
        "22. Hort M, Chen Z, Zhang JM, Harman M, Sarro F. Bias mitigation for machine learning classifiers: A comprehensive survey. ACM J Responsib Comput. 2024;1(2):1-52.",
        "23. Mhasawade V, Zhao Y, Chunara R. Machine learning and algorithmic fairness in public and population health. Nat Mach Intell. 2021;3(8):659-666.",
        "24. Regulation (EU) 2024/1689 of the European Parliament and of the Council of 13 June 2024 laying down harmonised rules on artificial intelligence (AI Act). OJ L, 2024/1689.",
        "25. Regulation (EU) 2016/679 of the European Parliament and of the Council of 27 April 2016 on the protection of natural persons with regard to the processing of personal data (GDPR). OJ L 119.",
        "26. Wachter S, Mittelstadt B, Russell C. Why fairness cannot be automated: Bridging the gap between EU non-discrimination law and AI. Comput Law Secur Rev. 2021;41:105567.",
        "27. Veale M, Binns R. Fairer machine learning in the real world: Mitigating discrimination without collecting sensitive data. Big Data Soc. 2017;4(2):1-17.",
        "28. European Parliament Research Service. Algorithmic discrimination under the AI Act and the GDPR. Brussels: EPRS; 2025.",
        "29. Deck C, Krafft PM, Stark P. It's complicated: The relationship of algorithmic fairness and non-discrimination regulations for high-risk systems in the EU AI Act. arXiv:2501.12962 [Preprint]. 2025.",
        "30. Prince MJ, Reischies F, Beekman AT, Fuhrer R, Jonker C, Kivela SL, et al. Development of the EURO-D scale\u2014a European Union initiative to compare symptoms of depression in 14 European centres. Br J Psychiatry. 1999;174(4):330-338.",
        "31. Hyde M, Wiggins RD, Higgs P, Blane DB. A measure of quality of life in early old age: The theory, development and properties of a needs satisfaction model (CASP-19). Aging Ment Health. 2003;7(3):186-194.",
        "32. Jylh\u00e4 M. What is self-rated health and why does it predict mortality? Towards a unified conceptual model. Soc Sci Med. 2009;68(3):553-560.",
        "33. Ke G, Meng Q, Finley T, Wang T, Chen W, Ma W, et al. LightGBM: A highly efficient gradient boosting decision tree. In: NeurIPS; 2017. pp. 3146-3154.",
        "34. Agarwal A, Beygelzimer A, Dud\u00edk M, Langford J, Wallach H. A reductions approach to fair classification. In: ICML; 2018. pp. 60-69.",
        "35. Cruz AF, Bel P, Soares C, Ghani R. FairGBM: Gradient boosting with fairness constraints. In: ICLR; 2023.",
        "36. Zafar MB, Valera I, Gomez Rodriguez M, Gummadi KP. Fairness constraints: Mechanisms for fair classification. J Mach Learn Res. 2019;20(75):1-42.",
        "37. Sagawa S, Koh PW, Hashimoto TB, Liang P. Distributionally robust neural networks for group shifts: On the importance of regularization for worst-case generalization. In: ICLR; 2020.",
        "38. Liu EZ, Haghgoo B, Chen AS, Raghunathan A, Koh PW, Sagawa S, et al. Just train twice: Improving group robustness without training group information. In: ICML; 2021. pp. 6781-6792.",
        "39. Dwork C, Immorlica N, Kalai AT, Leiserson M. Decoupled classifiers for group-fair and efficient machine learning. In: FAccT; 2018. pp. 119-133.",
        "40. Kearns M, Neel S, Roth A, Wu ZS. Preventing fairness gerrymandering: Auditing and learning for subgroup fairness. In: ICML; 2018. pp. 2564-2572.",
        "41. Kamiran F, Karim A, Zhang X. Decision theory for discrimination-aware classification. In: ICDM; 2012. pp. 924-929.",
        "42. Bellamy RKE, Dey K, Hind M, Hoffman SC, Houde S, Kannan K, et al. AI Fairness 360: An extensible toolkit for detecting and mitigating algorithmic bias. IBM J Res Dev. 2019;63(4/5):4:1-4:15.",
        "43. Bird S, Dud\u00edk M, Edgar R, Horn B, Lutz R, Milan V, et al. Fairlearn: A toolkit for assessing and improving fairness in AI. Microsoft Research; 2020.",
        "44. Chakraborty J, Majumder S, Menzies T. Bias in machine learning software: Why? How? What to do? In: ESEC/FSE; 2021. pp. 429-440.",
        "45. Xu L, Skoularidou M, Cuesta-Infante A, Veeramachaneni K. Modeling tabular data using conditional GAN. In: NeurIPS; 2019.",
        "46. Pleiss G, Raghavan M, Wu F, Kleinberg J, Weinberger KQ. On fairness and calibration. In: NeurIPS; 2017.",
        "47. Wyllie EC, Shumailov I, Papernot N. Fairness feedback loops: Training on synthetic data amplifies bias. In: FAccT; 2024.",
        "48. Montorsi C, Ferro Luzzi G, Ferretti C, Ferro Luzzi A. Predicting depression in old age: Combining life course data with machine learning. Econ Hum Biol. 2024;52:101330.",
        "49. Gum\u00e0 J, Arpino B. A machine learning approach to determine the influence of specific health conditions on self-rated health across education groups. BMC Public Health. 2023;23:197.",
        "50. Jin Y, Liu Y, Cai Y, Wang H, Shu H, Chen D, et al. FairMedFM: Fairness benchmarking for medical foundation models. In: NeurIPS Datasets and Benchmarks Track; 2024.",
        "51. BMC Digital Health. Post-processing methods for mitigating algorithmic bias in healthcare classification models: An extended umbrella review. BMC Digit Health. 2025;3:62.",
        "52. Stypi\u0144ska J, Franke A. AI ageism: A critical roadmap for studying age discrimination and exclusion in digitalized societies. Front Sociol. 2023;7:1017887.",
        "53. Luo Y, Carretta HJ, Li Q, Roth H, Prosperi M, Bian J. FAIM: Fairness-aware interpretable modeling for trustworthy machine learning in healthcare. Patterns. 2024;5(10):101061.",
        "54. Wang S, Guo W, Narasimhan H, Cotter A, Gupta M, Jordan MI. How robust is your fairness? Evaluating and sustaining fairness under unseen distribution shifts. Trans Mach Learn Res. 2023.",
        "55. Hardt M, Price E, Srebro N. Equality of opportunity in supervised learning. In: NeurIPS; 2016. pp. 3315-3323.",
        "56. Lundberg SM, Lee SI. A unified approach to interpreting model predictions. In: NeurIPS; 2017. pp. 4765-4774.",
        "57. Chen I, Johansson FD, Sontag D. Why is my classifier discriminatory? In: NeurIPS; 2018.",
        "58. Kamiran F, Calders T. Data preprocessing techniques for classification without discrimination. Knowl Inf Syst. 2012;33(1):1-33.",
        "59. Zhang BH, Lemoine B, Mitchell M. Mitigating unwanted biases with adversarial learning. In: AIES; 2018. pp. 335-340.",
        "60. Zemel R, Wu Y, Swersky K, Pitassi T, Dwork C. Learning fair representations. In: ICML; 2013. pp. 325-333.",
        "61. Feldman M, Friedler SA, Moeller J, Scheidegger C, Venkatasubramanian S. Certifying and removing disparate impact. In: KDD; 2015. pp. 259-268.",
        "62. Kamishima T, Akaho S, Asoh H, Sakuma J. Fairness-aware classifier with prejudice remover regularizer. In: ECML PKDD; 2012. pp. 35-50.",
        "63. Holstein K, Vaughan JW, Daum\u00e9 H III, Dud\u00edk M, Wallach H. Improving fairness in machine learning systems: What do industry practitioners need? In: CHI; 2019. Article 600.",
        "64. Chen RJ, Wang JJ, Williamson DFK, et al. Algorithmic fairness in artificial intelligence for medicine and healthcare. Nat Biomed Eng. 2023;7(6):719-742.",
        "65. Ferrara A, Langella A, Troiano L, Vitiello A. Fairness-aware machine learning engineering: How far are we? Empir Softw Eng. 2024;29(1):9.",
    ]
    for r in refs:
        p(r)

    # ================================================================
    # SUPPORTING INFORMATION
    # ================================================================
    h1("Supporting information")
    si = [
        "S1 Table. Complete benchmark results for all 26 methods across four health outcomes. AUROC, TPR gap, percentage reduction, 95% CI, equalised odds gap, demographic parity gap, AUROC gap, and sign test results.",
        "S2 Table. Multi-threshold robustness analysis. TPR gaps under four threshold strategies for all methods and outcomes.",
        "S3 Table. Cross-validation fold stability. Per-fold metrics and sign test results for all 26 methods.",
        "S4 Table. Computational timing. Training time per method per outcome.",
        "S5 Table. Variable definitions. All 175 predictor features with SHARE variable names and descriptions.",
        "S6 Table. Lifecycle filter results. Pass/fail status for all 26 methods at each lifecycle stage with survival outcome.",
        "S1 Fig. Pareto frontier plots. Fairness\u2013performance trade-off for all methods across four outcomes, with deployability status indicated.",
        "S2 Fig. ROC curves by income quintile. Baseline and top methods, stratified by quintile.",
        "S3 Fig. SHAP feature importance. Top 20 predictors for each outcome.",
        "S4 Fig. Predicted probability distributions by income group. Kernel density estimates for baseline and corrected models.",
        "S5 Fig. Fairea trade-off classification. Composite plot classifying each method\u2019s fairness improvement.",
        "S6 Fig. Lifecycle survival funnel. Progressive elimination of 26 methods through successive deployment stages.",
    ]
    for s in si:
        p(s)
        doc.add_paragraph()

    # ================================================================
    # FIGURE LEGENDS
    # ================================================================
    h1("Figure legends")
    figs = [
        "Fig 1. Cross-validation stability and benchmark overview. (A) Fold-level TPR gaps for the top 10 methods across four outcomes, demonstrating consistent improvement for constraint-based methods. (B) Percentage reduction in TPR gap versus AUROC change for all 26 methods; the upper-right quadrant indicates effective fairness at low accuracy cost. Deployable methods (Tier 1) are indicated by filled markers.",
        "Fig 2. Performance by method category with deployability annotation. Box plots showing TPR gap distributions for pre-processing, in-processing, post-processing, and augmentation methods across four outcomes. Deployability status is annotated: 0/6 post-processing methods are deployable versus 9/13 in-processing methods, illustrating the deployability cliff.",
        "Fig 3. Method heatmap. Percentage reduction in TPR gap for all 26 methods (rows) across four outcomes (columns). Green indicates improvement; red indicates that the method increased disparity. Reveals consistent top performers (ExpGrad variants) and widespread failure of augmentation and pre-processing approaches.",
        "Fig 4. Lifecycle survival funnel. Progressive elimination of 26 bias mitigation methods through successive lifecycle stages: effectiveness (26 \u2192 22), deployability (22 \u2192 11), accuracy preservation (11 \u2192 6), stability and calibration (6 \u2192 6), and fairness threshold (6 \u2192 3). Three methods survive the complete lifecycle: ExpGrad_TPR, ExpGrad_EO, and Reweighing.",
        "Fig 5. Clinical impact estimation. (A) Additional correct identifications per screening wave in the highest-income quintile (Q5) for each health outcome, comparing the three lifecycle survivors. (B) Total additional identifications across all four outcomes. ExpGrad_TPR yields approximately 1,723 additional correct identifications per screening wave.",
    ]
    for f in figs:
        p(f)
        doc.add_paragraph()

    doc.save("plos_manuscript_v3.docx")
    print("Saved: plos_manuscript_v3.docx")


if __name__ == "__main__":
    build()
