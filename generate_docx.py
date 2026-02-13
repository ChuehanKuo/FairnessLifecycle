"""Convert the narrative markdown to a formatted Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn('w:shd'),
        {qn('w:fill'): color_hex, qn('w:val'): 'clear'}
    )
    shading.append(shading_elm)


def add_formatted_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s:-]+\|$', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, cell_text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text.replace('**', ''))
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, '2E4057')
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row_data in enumerate(rows[1:], 1):
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean = cell_text.replace('**', '').replace('*', '')
            run = p.add_run(clean)
            run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'F0F4F8')

    doc.add_paragraph()


def process_inline(paragraph, text):
    """Process bold and italic markers in text."""
    # Split by bold markers first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            # Check for italic inside bold
            italic_parts = re.split(r'(\*.*?\*)', inner)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    add_formatted_run(paragraph, ip[1:-1], bold=True, italic=True)
                else:
                    add_formatted_run(paragraph, ip, bold=True)
        else:
            # Check for italic
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    add_formatted_run(paragraph, ip[1:-1], italic=True)
                else:
                    paragraph.add_run(ip)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open('NARRATIVE_AND_DEPLOYMENT_SIMULATION.md', 'r') as f:
        content = f.read()

    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            doc.add_paragraph()
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # H1 Title
        if stripped.startswith('# ') and not stripped.startswith('##'):
            title = stripped[2:]
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2
        if stripped.startswith('## ') and not stripped.startswith('###'):
            heading_text = stripped[3:]
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # H3
        if stripped.startswith('### ') and not stripped.startswith('####'):
            heading_text = stripped[4:]
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        # H4
        if stripped.startswith('#### '):
            heading_text = stripped[5:]
            p = doc.add_paragraph()
            add_formatted_run(p, heading_text.replace('**', ''), bold=True,
                              size=Pt(12), color=RGBColor(0x2E, 0x40, 0x57))
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Inches(0.3)
            for ci, cl in enumerate(code_lines):
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
                if ci < len(code_lines) - 1:
                    p.add_run('\n')
            continue

        # Block quote
        if stripped.startswith('> '):
            quote_text = stripped[2:].replace('> ', '')
            # Collect multi-line quotes
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                quote_text += ' ' + lines[i].strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(quote_text.replace('"', '\u201c', 1).replace('"', '\u201d', 1))
            run.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tl = lines[i].strip()
                # Skip separator lines
                if not re.match(r'^\|[\s:-]+\|$', tl):
                    table_lines.append(tl)
                i += 1
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            continue

        # Bullet list
        if stripped.startswith('- '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, bullet_text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            list_text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, list_text)
            i += 1
            continue

        # Regular paragraph - collect continuation lines
        para_text = stripped
        while (i + 1 < len(lines) and
               lines[i + 1].strip() and
               not lines[i + 1].strip().startswith(('#', '|', '-', '>', '```', '---')) and
               not re.match(r'^\d+\.', lines[i + 1].strip())):
            i += 1
            para_text += ' ' + lines[i].strip()

        p = doc.add_paragraph()
        process_inline(p, para_text)
        i += 1

    output_path = 'NARRATIVE_AND_DEPLOYMENT_SIMULATION.docx'
    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    main()
