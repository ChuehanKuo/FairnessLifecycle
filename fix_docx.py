"""Fix structural issues in FairnessLifecycleManuscript_corrected.docx:
1. Number all 78 references
2. Embed Tables 1-3 as Word table objects
"""
import csv
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell_text(cell, text, bold=False, size=8, font_name="Calibri"):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold


def set_cell_left(cell, text, bold=False, size=8, font_name="Calibri"):
    """Set cell text left-aligned."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold


def shade_cells(row, color_hex):
    """Apply shading to all cells in a row."""
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.find(qn("w:shd"))
        if shd is None:
            shd = shading.makeelement(qn("w:shd"), {})
            shading.append(shd)
        shd.set(qn("w:fill"), color_hex)
        shd.set(qn("w:val"), "clear")


def set_table_borders(table):
    """Set thin borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tblPr.makeelement(qn("w:tblBorders"), {})
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "999999",
        })
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        borders.append(elem)


def set_col_width(table, col_idx, width_cm):
    """Set column width."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def load_csv(path):
    with open(path, newline="") as f:
        return list(csv.DictReader(f))


def build_table1(doc, insert_after_para_idx):
    """Build Table 1: Baseline model performance and income-based disparities."""
    rows_data = load_csv("/home/user/FairnessLifecycle/tables/table1_descriptive.csv")
    full = load_csv("/home/user/FairnessLifecycle/tables/full_results.csv")

    # Get baseline rows
    baselines = {r["Outcome"]: r for r in full if r["Method"] == "BASELINE"}

    # Table 1 has two panels:
    # Panel A: Prevalence by quintile (from table1_descriptive.csv)
    # Panel B: Baseline model performance (AUROC, TPR gap)

    # Create combined table
    # Columns: SES Quintile | N | EUROD N(%) | LS N(%) | CASP N(%) | SRH N(%)
    headers = ["SES Quintile", "N", "EUROD N(%)", "LS N(%)", "CASP N(%)", "SRH N(%)"]

    # Panel A: 1 header + 6 data rows + 1 blank + Panel B header + 4 data rows = ~13 rows
    n_rows = 1 + len(rows_data) + 1 + 1 + 1  # header + data + blank + panel B header + baseline metrics
    table = doc.add_table(rows=0, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Panel A header
    row = table.add_row()
    shade_cells(row, "D9E2F3")
    for i, h in enumerate(["Income Quintile", "N", "Depression\nN (%)", "Life Satisfaction\nN (%)", "Quality of Life\nN (%)", "Self-Rated Health\nN (%)"]):
        set_cell_text(row.cells[i], h, bold=True, size=8)

    # Panel A data
    for rd in rows_data:
        row = table.add_row()
        set_cell_left(row.cells[0], rd["SES Quintile"], bold=(rd["SES Quintile"] == "Total"), size=8)
        set_cell_text(row.cells[1], rd["N"], size=8)
        set_cell_text(row.cells[2], rd["EUROD N(%)"], size=8)
        set_cell_text(row.cells[3], rd["LS N(%)"], size=8)
        set_cell_text(row.cells[4], rd["CASP N(%)"], size=8)
        set_cell_text(row.cells[5], rd["SRH N(%)"], size=8)
        if rd["SES Quintile"] == "Total":
            shade_cells(row, "F2F2F2")

    # Blank separator row
    row = table.add_row()
    for c in row.cells:
        set_cell_text(c, "", size=4)

    # Panel B: Baseline model performance
    row = table.add_row()
    shade_cells(row, "D9E2F3")
    set_cell_left(row.cells[0], "Baseline Model", bold=True, size=8)
    for i, outcome in enumerate(["", "EUROD", "LS", "CASP", "SRH"]):
        if i > 0:
            set_cell_text(row.cells[i+1], outcome, bold=True, size=8)
    set_cell_text(row.cells[1], "", size=8)

    # AUROC row
    row = table.add_row()
    set_cell_left(row.cells[0], "AUROC", size=8)
    set_cell_text(row.cells[1], "", size=8)
    for i, outcome in enumerate(["EUROD", "LS", "CASP", "SRH"]):
        b = baselines[outcome]
        auroc = f"{float(b['AUROC_mean']):.3f}"
        set_cell_text(row.cells[i+2], auroc, size=8)

    # TPR Gap row
    row = table.add_row()
    set_cell_left(row.cells[0], "TPR Gap (Q1 vs Q5)", size=8)
    set_cell_text(row.cells[1], "", size=8)
    for i, outcome in enumerate(["EUROD", "LS", "CASP", "SRH"]):
        b = baselines[outcome]
        gap = f"{float(b['TPR_gap_mean']):.3f}"
        set_cell_text(row.cells[i+2], gap, size=8)

    # Sensitivity row
    row = table.add_row()
    set_cell_left(row.cells[0], "Sensitivity", size=8)
    set_cell_text(row.cells[1], "", size=8)
    for i, outcome in enumerate(["EUROD", "LS", "CASP", "SRH"]):
        b = baselines[outcome]
        sens = f"{float(b['Sens_mean']):.3f}"
        set_cell_text(row.cells[i+2], sens, size=8)

    # Specificity row
    row = table.add_row()
    set_cell_left(row.cells[0], "Specificity", size=8)
    set_cell_text(row.cells[1], "", size=8)
    for i, outcome in enumerate(["EUROD", "LS", "CASP", "SRH"]):
        b = baselines[outcome]
        spec = f"{float(b['Spec_mean']):.3f}"
        set_cell_text(row.cells[i+2], spec, size=8)

    # Move table to after the caption paragraph
    para_elem = doc.paragraphs[insert_after_para_idx]._element
    para_elem.addnext(table._tbl)

    return table


def build_table2(doc, insert_after_para_idx):
    """Build Table 2: Benchmark performance of 26 methods (summary across 4 outcomes)."""
    full = load_csv("/home/user/FairnessLifecycle/tables/full_results.csv")

    # Group by method, compute mean across outcomes
    methods = {}
    for r in full:
        if r["Method"] == "BASELINE":
            continue
        m = r["Method"]
        if m not in methods:
            methods[m] = {"category": r["Category"], "outcomes": {}}
        methods[m]["outcomes"][r["Outcome"]] = r

    # Build rows: Method | Category | AUROC range | AUROC change range | TPR Gap range | % Reduction range | Deployable
    table = doc.add_table(rows=0, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header
    row = table.add_row()
    shade_cells(row, "D9E2F3")
    hdrs = ["Method", "Category", "AUROC\n(range)", "\u0394AUROC\n(pp)", "TPR Gap\n(range)", "% Reduction\n(range)", "Deploy."]
    for i, h in enumerate(hdrs):
        set_cell_text(row.cells[i], h, bold=True, size=7)

    # Sort by category then by mean % reduction descending
    def sort_key(item):
        cat_order = {"Pre-processing": 0, "In-processing": 1, "Post-processing": 2, "Augmentation": 3}
        m, info = item
        reds = [float(info["outcomes"][o]["Pct_Reduction"]) for o in info["outcomes"]]
        return (cat_order.get(info["category"], 4), -sum(reds)/len(reds))

    sorted_methods = sorted(methods.items(), key=sort_key)

    # Get baselines for AUROC change calc
    baselines = {r["Outcome"]: float(r["AUROC_mean"]) for r in full if r["Method"] == "BASELINE"}

    prev_cat = None
    for method_name, info in sorted_methods:
        outs = info["outcomes"]
        aurocs = [float(outs[o]["AUROC_mean"]) for o in outs]
        auroc_changes = [(float(outs[o]["AUROC_mean"]) - baselines[o]) * 100 for o in outs]
        tpr_gaps = [float(outs[o]["TPR_gap_mean"]) for o in outs]
        pct_reds = [float(outs[o]["Pct_Reduction"]) for o in outs]

        deployable_data = load_csv("/home/user/FairnessLifecycle/tables/deployability_contract.csv")
        deploy_map = {r["Method"]: r["Deployable"] for r in deployable_data}
        is_deploy = deploy_map.get(method_name, "?")

        row = table.add_row()

        # Category separator shading
        if info["category"] != prev_cat and prev_cat is not None:
            shade_cells(row, "F7F7F7")

        set_cell_left(row.cells[0], method_name, size=7)
        set_cell_text(row.cells[1], info["category"].replace("-processing", "-proc.").replace("Augmentation", "Augment."), size=7)
        set_cell_text(row.cells[2], f"{min(aurocs):.3f}\u2013{max(aurocs):.3f}", size=7)

        min_ch, max_ch = min(auroc_changes), max(auroc_changes)
        set_cell_text(row.cells[3], f"{min_ch:+.1f} to {max_ch:+.1f}", size=7)
        set_cell_text(row.cells[4], f"{min(tpr_gaps):.3f}\u2013{max(tpr_gaps):.3f}", size=7)
        set_cell_text(row.cells[5], f"{min(pct_reds):.1f}\u2013{max(pct_reds):.1f}", size=7)
        set_cell_text(row.cells[6], "\u2713" if is_deploy == "Yes" else "\u2717", size=7)

        prev_cat = info["category"]

    # Move table after caption
    para_elem = doc.paragraphs[insert_after_para_idx]._element
    para_elem.addnext(table._tbl)

    return table


def build_table3(doc, insert_after_para_idx):
    """Build Table 3: Deployability assessment."""
    deploy_data = load_csv("/home/user/FairnessLifecycle/tables/deployability_contract.csv")

    table = doc.add_table(rows=0, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header
    row = table.add_row()
    shade_cells(row, "D9E2F3")
    hdrs = ["Method", "Category", "A at\nTraining", "A at\nInference", "A for\nThreshold", "Deployable"]
    for i, h in enumerate(hdrs):
        set_cell_text(row.cells[i], h, bold=True, size=8)

    # Sort by category
    cat_order = {"Pre-processing": 0, "In-processing": 1, "Post-processing": 2, "Augmentation": 3}
    deploy_data.sort(key=lambda r: (cat_order.get(r["Category"], 4), r["Method"]))

    prev_cat = None
    for r in deploy_data:
        row = table.add_row()

        set_cell_left(row.cells[0], r["Method"], size=8)
        set_cell_text(row.cells[1], r["Category"], size=8)

        for i, field in enumerate(["A_at_train", "A_at_inference", "A_for_threshold"]):
            val = r[field]
            sym = "\u2713" if val == "Yes" else "\u2717"
            set_cell_text(row.cells[i+2], sym, size=8)

        deploy_sym = "\u2713" if r["Deployable"] == "Yes" else "\u2717"
        set_cell_text(row.cells[5], deploy_sym, bold=True, size=8)

        # Red shading for non-deployable
        if r["Deployable"] == "No":
            shade_cells(row, "FDE8E8")

        prev_cat = r["Category"]

    para_elem = doc.paragraphs[insert_after_para_idx]._element
    para_elem.addnext(table._tbl)

    return table


def fix_references(doc):
    """Number all references in the References section."""
    in_refs = False
    ref_num = 0
    ref_start = None

    for i, para in enumerate(doc.paragraphs):
        if "Heading" in para.style.name and para.text.strip().lower() == "references":
            in_refs = True
            ref_start = i
            continue
        if in_refs and "Heading" in para.style.name:
            break
        if in_refs and para.text.strip():
            ref_num += 1
            # Prepend number to the text
            text = para.text.strip()
            # Check if already numbered
            if not text[0].isdigit():
                # Clear and re-add with number
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = f"{ref_num}. {text}"

    print(f"Numbered {ref_num} references")


def main():
    doc = Document("/home/user/FairnessLifecycle/FairnessLifecycleManuscript_corrected.docx")

    # 1. Fix references
    fix_references(doc)

    # 2. Find table caption paragraph indices
    # Table captions are at paragraphs [48], [51], [56] based on analysis
    table_captions = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("Table 1."):
            table_captions[1] = i
        elif text.startswith("Table 2."):
            table_captions[2] = i
        elif text.startswith("Table 3."):
            table_captions[3] = i

    print(f"Table caption indices: {table_captions}")

    # Insert tables AFTER their captions (insert in reverse order to preserve indices)
    # Table 3 first (highest index), then 2, then 1
    build_table3(doc, table_captions[3])
    build_table2(doc, table_captions[2])
    build_table1(doc, table_captions[1])

    # Save
    out_path = "/home/user/FairnessLifecycle/FairnessLifecycleManuscript_corrected.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
